#!/usr/bin/env python3
"""Build the Latticra Reference Manual artifacts.

The manual is generated from the repository's documentation sources so future
release updates can rebuild the PDF, DOCX, editable Markdown, and manifest from
one repeatable command.
"""

from __future__ import annotations

import argparse
import hashlib
import html
import json
import re
import textwrap
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT_FRONT_MATTER = [
    "README.md",
    "STATUS.md",
    "SECURITY.md",
    "CONTRIBUTING.md",
    "TRADEMARK_POLICY.md",
]

OPENING_DOCS = [
    "docs/README.md",
    "docs/QUICK_START_CHEATSHEET.md",
    "docs/DOCUMENTATION_READER_JOURNEY_MAP.md",
    "docs/DOCUMENTATION_GLOSSARY.md",
    "docs/PRODUCT_DOCUMENTATION_COHESION.md",
    "docs/PUBLIC_CLAIMS_LEDGER.md",
    "docs/NON_CLAIMS.md",
    "docs/REAL_SYSTEM_CONTRACT.md",
    "docs/EVIDENCE_LADDER.md",
    "docs/PRECURSOR_PROMOTION_RULE.md",
    "docs/FOUNDATION_INDEX.md",
]

DOC_SUFFIXES = (".md", ".txt")

EXCLUDED_DIR_PARTS = {
    ".git",
    "build",
    "target",
    "__pycache__",
    "latticra-reference-manual",
}

EXCLUDED_PREFIXES = (
    "reports/",
    "installer/latticra-installer-receipts/",
)

INCLUDED_TEXT_FILES = {
    "LICENSE",
    "installer/TREE.txt",
    "installer/latticra-installer-plan.txt",
    "fixtures/artifact/local-artifact-manifest.txt",
}


@dataclass(frozen=True)
class SourceDoc:
    path: str
    title: str
    part: str
    lines: int
    words: int
    bytes: int
    sha256: str
    content: str


PART_ORDER = [
    "Part I - Orientation and Reader Routes",
    "Part II - Current Status, Evidence, and Public Boundaries",
    "Part III - Architecture and Foundation Concepts",
    "Part IV - Lat, LIR, and L-UI",
    "Part V - Nucleus and Runtime Boundary",
    "Part VI - Latticra Seal",
    "Part VII - Console, Panel, Installer, and Update Paths",
    "Part VIII - Nadia Offline AI Foundation",
    "Part IX - Kernel, Boot, and OS Image Research",
    "Part X - Security and Assurance",
    "Part XI - Platform Packaging and Distribution Lanes",
    "Part XII - Project Notes, Strategy, and Documentation Maintenance",
    "Part XIII - Status Record Archive",
    "Part XIV - Legal, License, and Remaining Records",
]


def rel(path: Path, root: Path) -> str:
    return path.relative_to(root).as_posix()


def should_include(path: Path, root: Path) -> bool:
    rel_path = rel(path, root)
    parts = set(path.relative_to(root).parts)

    if parts & EXCLUDED_DIR_PARTS:
        return False
    if any(rel_path.startswith(prefix) for prefix in EXCLUDED_PREFIXES):
        return False
    if rel_path in INCLUDED_TEXT_FILES:
        return True
    if rel_path.startswith("LICENSES/") and path.suffix == ".txt":
        return True
    if path.suffix not in DOC_SUFFIXES:
        return False
    if rel_path.startswith("docs/"):
        return True
    if rel_path.startswith("installer/") and path.suffix in DOC_SUFFIXES:
        return True
    if rel_path.startswith("packaging/") and path.name == "README.md":
        return True
    if rel_path in ROOT_FRONT_MATTER or rel_path == "LICENSE":
        return True
    if rel_path.startswith("LICENSES/") and path.name == "README.md":
        return True
    return False


def extract_title(path: str, text: str) -> str:
    in_fence = False
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("```") or stripped.startswith("~~~"):
            in_fence = not in_fence
        if in_fence:
            continue
        match = re.match(r"^#{1,6}\s+(.+?)\s*$", stripped)
        if match:
            return clean_inline(match.group(1))
    stem = Path(path).stem.replace("_", " ").replace("-", " ")
    return stem.title()


def clean_inline(value: str) -> str:
    value = re.sub(r"`([^`]+)`", r"\1", value)
    value = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", value)
    value = value.replace("**", "").replace("__", "").replace("*", "")
    return value.strip()


def classify(path: str) -> str:
    name = Path(path).name
    upper = name.upper()

    if path in ROOT_FRONT_MATTER or path in OPENING_DOCS:
        return "Part I - Orientation and Reader Routes"
    if path.startswith("docs/status/"):
        if name in {"CURRENT_STATUS.md", "README.md", "ANNOUNCEMENTS.md"}:
            return "Part II - Current Status, Evidence, and Public Boundaries"
        return "Part XIII - Status Record Archive"
    if upper in {
        "REAL_SYSTEM_CONTRACT.MD",
        "EVIDENCE_LADDER.MD",
        "PRECURSOR_PROMOTION_RULE.MD",
        "FOUNDATION_INDEX.MD",
        "PUBLIC_CLAIMS_LEDGER.MD",
        "NON_CLAIMS.MD",
        "ROADMAP.MD",
        "LATTICRA_LAT_SEAL_ROADMAP.MD",
    }:
        return "Part II - Current Status, Evidence, and Public Boundaries"
    if path.startswith("docs/latticra-seal/") or upper.startswith("LATTICRA_SEAL"):
        return "Part VI - Latticra Seal"
    if path.startswith("docs/specs/"):
        return "Part VI - Latticra Seal"
    if upper.startswith("NADIA_") or "NADIA" in upper:
        return "Part VIII - Nadia Offline AI Foundation"
    if path.startswith("installer/") or "INSTALLER" in upper or "PANEL" in upper or "UPDATER" in upper:
        return "Part VII - Console, Panel, Installer, and Update Paths"
    if "LATTICRA_CONSOLE" in upper or "SELF_UPDATE" in upper:
        return "Part VII - Console, Panel, Installer, and Update Paths"
    if upper.startswith(("LAT_", "LIR_", "L_UI_")) or upper.startswith("UI_TERMINAL"):
        return "Part IV - Lat, LIR, and L-UI"
    if "LANGUAGE" in upper and not path.startswith("docs/status/"):
        return "Part IV - Lat, LIR, and L-UI"
    if upper.startswith(("NUCLEUS", "RUNTIME_BOUNDARY", "TRI_PLANE", "STATE_LATTICE")):
        return "Part V - Nucleus and Runtime Boundary"
    if upper.startswith(("KERNEL", "SEABIOS")) or "OS_IMAGE" in upper or "BOOT" in upper:
        return "Part IX - Kernel, Boot, and OS Image Research"
    if path.startswith("docs/security/") or any(
        token in upper
        for token in (
            "SECURITY",
            "ASSURANCE",
            "THREAT",
            "ZERO_TRUST",
            "SUPPLY_CHAIN",
            "VULNERABILITY",
            "CRYPTOGRAPHIC",
            "IDENTITY",
            "BACKUP",
            "INCIDENT",
            "MEMORY_SAFETY",
            "SECURE_CONFIGURATION",
        )
    ):
        return "Part X - Security and Assurance"
    if path.startswith("packaging/") or any(
        token in upper
        for token in (
            "FEDORA",
            "UBUNTU",
            "DEBIAN",
            "OPENSUSE",
            "FREEBSD",
            "OPENBSD",
            "PACKAGE",
            "RPM",
            "DEB",
            "PORT",
            "PPA",
            "LINTIAN",
            "RPMLINT",
            "SOURCE_ARCHIVE",
        )
    ):
        return "Part XI - Platform Packaging and Distribution Lanes"
    if path.startswith("docs/project_notes/") or path.startswith("docs/strategy/"):
        return "Part XII - Project Notes, Strategy, and Documentation Maintenance"
    if "DOCUMENTATION_" in upper or "LICENSE_" in upper:
        return "Part XII - Project Notes, Strategy, and Documentation Maintenance"
    if any(
        token in upper
        for token in (
            "ARCHITECTURE",
            "FOUNDATION",
            "C_CPP",
            "CPP",
            "AUTHORITY",
            "SUPERVISOR",
            "EFFECT",
            "SERVER",
            "OPEN_ECOSYSTEM",
            "NAMING",
            "FEATURE",
        )
    ):
        return "Part III - Architecture and Foundation Concepts"
    return "Part XIV - Legal, License, and Remaining Records"


def discover_sources(root: Path) -> list[SourceDoc]:
    docs: list[SourceDoc] = []
    for path in root.rglob("*"):
        if not path.is_file():
            continue
        if not should_include(path, root):
            continue
        rel_path = rel(path, root)
        text = path.read_text(encoding="utf-8", errors="replace")
        encoded = text.encode("utf-8", errors="replace")
        docs.append(
            SourceDoc(
                path=rel_path,
                title=extract_title(rel_path, text),
                part=classify(rel_path),
                lines=len(text.splitlines()),
                words=len(re.findall(r"\S+", text)),
                bytes=len(encoded),
                sha256=hashlib.sha256(encoded).hexdigest(),
                content=text,
            )
        )

    part_index = {part: index for index, part in enumerate(PART_ORDER)}
    front_order = {path: index for index, path in enumerate(ROOT_FRONT_MATTER + OPENING_DOCS)}

    def sort_key(doc: SourceDoc) -> tuple[int, int, str]:
        return (
            part_index.get(doc.part, 999),
            front_order.get(doc.path, 999),
            doc.path.lower(),
        )

    return sorted(docs, key=sort_key)


def companion_artifacts(root: Path) -> dict[str, list[str]]:
    html_files = sorted(rel(p, root) for p in (root / "docs").glob("*.html"))
    pdf_files = sorted(rel(p, root) for p in root.rglob("*.pdf") if "latticra-reference-manual" not in p.parts)
    docx_files = sorted(rel(p, root) for p in root.rglob("*.docx") if "latticra-reference-manual" not in p.parts)
    return {
        "public_html_surfaces_indexed_not_inlined": html_files,
        "existing_pdf_artifacts_indexed_not_inlined": pdf_files,
        "existing_docx_artifacts_indexed_not_inlined": docx_files,
    }


def heading_shifted_markdown(text: str, offset: int = 2) -> str:
    output: list[str] = []
    in_fence = False
    fence_token = ""
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith(("```", "~~~")):
            token = stripped[:3]
            if not in_fence:
                in_fence = True
                fence_token = token
            elif token == fence_token:
                in_fence = False
                fence_token = ""
            output.append(line.rstrip())
            continue
        if not in_fence:
            match = re.match(r"^(#{1,6})(\s+.+)$", line)
            if match:
                level = min(6, len(match.group(1)) + offset)
                output.append("#" * level + match.group(2))
                continue
        output.append(line.rstrip())
    return "\n".join(output).strip()


def build_markdown(
    docs: Sequence[SourceDoc],
    artifacts: dict[str, list[str]],
    output_path: Path,
    generated_at: str,
) -> None:
    totals = {
        "documents": len(docs),
        "lines": sum(doc.lines for doc in docs),
        "words": sum(doc.words for doc in docs),
        "bytes": sum(doc.bytes for doc in docs),
    }
    lines: list[str] = [
        "# The Latticra Reference Manual",
        "",
        "**Edition:** v0.1.0 documentation assembly",
        f"**Generated:** {generated_at}",
        "**Scope:** Canonical repository documentation sources in the current working tree.",
        "",
        "This manual collects the Latticra documentation corpus into one ordered reference book. "
        "Generated public HTML pages and existing PDF/DOCX handbooks are indexed as companion artifacts rather than inlined, because their source records are already represented in the canonical Markdown and text corpus.",
        "",
        "## Compilation Boundary",
        "",
        "- The manual preserves documentation source content while adding a reference-manual structure around it.",
        "- The source snapshot follows the current working tree, including uncommitted documentation changes.",
        "- Runtime, package, security, OS, and production-readiness claims remain bounded by the source records and non-claims included below.",
        "- The root README links directly to this manual package so the PDF, editable source, and source manifest are easy to find.",
        "",
        "## Source Corpus Summary",
        "",
        f"- Source documents: {totals['documents']}",
        f"- Source lines: {totals['lines']}",
        f"- Source words: {totals['words']}",
        f"- Source bytes: {totals['bytes']}",
        "",
        "## Manual Part Order",
        "",
    ]

    docs_by_part = group_by_part(docs)
    for part in PART_ORDER:
        part_docs = docs_by_part.get(part, [])
        if part_docs:
            lines.append(f"- {part}: {len(part_docs)} source records")

    lines.extend(["", "## Companion Artifact Index", ""])
    for label, values in artifacts.items():
        readable = label.replace("_", " ").title()
        lines.append(f"### {readable}")
        lines.append("")
        if values:
            for value in values:
                lines.append(f"- `{value}`")
        else:
            lines.append("- None found.")
        lines.append("")

    lines.extend(["# Source Corpus", ""])
    for part in PART_ORDER:
        part_docs = docs_by_part.get(part, [])
        if not part_docs:
            continue
        lines.extend([f"# {part}", ""])
        lines.append("## Included Source Records")
        lines.append("")
        for doc in part_docs:
            lines.append(f"- `{doc.path}` - {doc.title}")
        lines.append("")
        for doc in part_docs:
            lines.extend(
                [
                    f"## {doc.path}",
                    "",
                    f"**Source title:** {doc.title}",
                    f"**Source path:** `{doc.path}`",
                    f"**Lines:** {doc.lines}",
                    f"**Words:** {doc.words}",
                    f"**SHA-256:** `{doc.sha256}`",
                    "",
                    "<!-- BEGIN SOURCE DOCUMENT -->",
                    "",
                    heading_shifted_markdown(doc.content),
                    "",
                    "<!-- END SOURCE DOCUMENT -->",
                    "",
                ]
            )

    output_path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


def group_by_part(docs: Sequence[SourceDoc]) -> dict[str, list[SourceDoc]]:
    grouped: dict[str, list[SourceDoc]] = {part: [] for part in PART_ORDER}
    for doc in docs:
        grouped.setdefault(doc.part, []).append(doc)
    return grouped


def write_manifest(
    docs: Sequence[SourceDoc],
    artifacts: dict[str, list[str]],
    output_path: Path,
    generated_at: str,
) -> None:
    manifest = {
        "title": "The Latticra Reference Manual",
        "edition": "v0.1.0 documentation assembly",
        "generated_at": generated_at,
        "source_count": len(docs),
        "total_lines": sum(doc.lines for doc in docs),
        "total_words": sum(doc.words for doc in docs),
        "total_bytes": sum(doc.bytes for doc in docs),
        "manual_order": PART_ORDER,
        "sources": [
            {
                "path": doc.path,
                "title": doc.title,
                "part": doc.part,
                "lines": doc.lines,
                "words": doc.words,
                "bytes": doc.bytes,
                "sha256": doc.sha256,
            }
            for doc in docs
        ],
        "companion_artifacts": artifacts,
        "readme_link_policy": "Root README links to the manual package from the top handbook links, Start Here, and Main Documentation.",
    }
    output_path.write_text(json.dumps(manifest, indent=2) + "\n", encoding="utf-8")


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold


def configure_docx_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    for name, size, color in [
        ("Title", 24, "0B2545"),
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
        ("Heading 4", 11, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    if "ReferenceCode" not in styles:
        code_style = styles.add_style("ReferenceCode", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = styles["ReferenceCode"]
    code_style.font.name = "Courier New"
    code_style.font.size = Pt(8)
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    code_style.paragraph_format.space_after = Pt(0)
    code_style.paragraph_format.line_spacing = 1.0

    if "ReferenceMeta" not in styles:
        meta_style = styles.add_style("ReferenceMeta", WD_STYLE_TYPE.PARAGRAPH)
    else:
        meta_style = styles["ReferenceMeta"]
    meta_style.font.name = "Calibri"
    meta_style.font.size = Pt(9)
    meta_style.font.color.rgb = RGBColor.from_string("555555")
    meta_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")


def add_docx_field(paragraph, field: str) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def add_docx_footer(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.style = "ReferenceMeta"
    paragraph.add_run("The Latticra Reference Manual - ")
    add_docx_field(paragraph, "PAGE")


def build_docx(docs: Sequence[SourceDoc], output_path: Path, generated_at: str) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    configure_docx_styles(document)
    add_docx_footer(section)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("The Latticra Reference Manual").bold = True

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("v0.1.0 Documentation Assembly").italic = True
    document.add_paragraph(f"Generated: {generated_at}", style="ReferenceMeta").alignment = WD_ALIGN_PARAGRAPH.CENTER

    summary = document.add_table(rows=5, cols=2)
    summary.style = "Table Grid"
    rows = [
        ("Source documents", str(len(docs))),
        ("Source lines", str(sum(doc.lines for doc in docs))),
        ("Source words", str(sum(doc.words for doc in docs))),
        ("Editable source", "the-latticra-reference-manual-v0.1.0.md"),
        ("Root README link", "Top handbook links + Start Here + Main Documentation"),
    ]
    for row, (label, value) in zip(summary.rows, rows):
        set_cell_text(row.cells[0], label, bold=True)
        set_cell_text(row.cells[1], value)

    document.add_paragraph(
        "This generated reference book preserves the current working-tree documentation corpus in a reader-oriented order. "
        "Public HTML pages and existing companion PDFs/DOCX files are indexed separately in the manifest.",
    )

    document.add_page_break()
    document.add_heading("Manual Part Order", level=1)
    docs_by_part = group_by_part(docs)
    for part in PART_ORDER:
        count = len(docs_by_part.get(part, []))
        if count:
            document.add_paragraph(f"{part}: {count} source records", style="List Bullet")

    for part in PART_ORDER:
        part_docs = docs_by_part.get(part, [])
        if not part_docs:
            continue
        document.add_page_break()
        document.add_heading(part, level=1)
        for doc in part_docs:
            document.add_paragraph(f"{doc.path} - {doc.title}", style="List Bullet")
        for doc in part_docs:
            document.add_heading(doc.path, level=2)
            document.add_paragraph(f"Source title: {doc.title}", style="ReferenceMeta")
            document.add_paragraph(
                f"Lines: {doc.lines} | Words: {doc.words} | SHA-256: {doc.sha256}",
                style="ReferenceMeta",
            )
            add_markdown_to_docx(document, doc.content)

    document.save(output_path)


def flush_paragraph_docx(document: Document, lines: list[str]) -> None:
    if not lines:
        return
    text = " ".join(line.strip() for line in lines if line.strip())
    if text:
        document.add_paragraph(clean_inline(text))
    lines.clear()


def add_markdown_to_docx(document: Document, markdown_text: str) -> None:
    paragraph_lines: list[str] = []
    code_lines: list[str] = []
    in_code = False
    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip("\n")
        stripped = line.strip()
        if stripped.startswith(("```", "~~~")):
            flush_paragraph_docx(document, paragraph_lines)
            if in_code:
                add_code_block_docx(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not stripped:
            flush_paragraph_docx(document, paragraph_lines)
            continue
        if is_markdown_table_line(stripped):
            flush_paragraph_docx(document, paragraph_lines)
            document.add_paragraph(line, style="ReferenceCode")
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            flush_paragraph_docx(document, paragraph_lines)
            level = min(4, len(heading.group(1)) + 2)
            document.add_heading(clean_inline(heading.group(2)), level=level)
            continue
        bullet = re.match(r"^[-*+]\s+(.+)$", stripped)
        if bullet:
            flush_paragraph_docx(document, paragraph_lines)
            document.add_paragraph(clean_inline(bullet.group(1)), style="List Bullet")
            continue
        number = re.match(r"^\d+[.)]\s+(.+)$", stripped)
        if number:
            flush_paragraph_docx(document, paragraph_lines)
            document.add_paragraph(clean_inline(number.group(1)), style="List Number")
            continue
        if stripped.startswith(">"):
            flush_paragraph_docx(document, paragraph_lines)
            document.add_paragraph(clean_inline(stripped.lstrip("> ")))
            continue
        if stripped.startswith("|") or stripped.startswith("---"):
            flush_paragraph_docx(document, paragraph_lines)
            document.add_paragraph(line, style="ReferenceCode")
            continue
        paragraph_lines.append(line)
    flush_paragraph_docx(document, paragraph_lines)
    if code_lines:
        add_code_block_docx(document, code_lines)


def add_code_block_docx(document: Document, lines: Sequence[str]) -> None:
    if not lines:
        return
    for line in lines:
        chunks = wrap_preserving_words(line, 110) or [""]
        for chunk in chunks:
            document.add_paragraph(chunk, style="ReferenceCode")


def is_markdown_table_line(line: str) -> bool:
    return line.startswith("|") and line.endswith("|")


def wrap_preserving_words(text: str, width: int) -> list[str]:
    if len(text) <= width:
        return [text]
    return textwrap.wrap(
        text,
        width=width,
        replace_whitespace=False,
        drop_whitespace=False,
        break_long_words=True,
        break_on_hyphens=False,
    )


def pdf_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "ManualTitle",
            parent=base["Title"],
            fontName="Courier-Bold",
            fontSize=24,
            leading=30,
            textColor=colors.HexColor("#0B2545"),
            alignment=TA_CENTER,
            spaceAfter=12,
        ),
        "subtitle": ParagraphStyle(
            "ManualSubtitle",
            parent=base["Normal"],
            fontName="Courier-Oblique",
            fontSize=12,
            leading=15,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#555555"),
            spaceAfter=18,
        ),
        "h1": ParagraphStyle(
            "ManualH1",
            parent=base["Heading1"],
            fontName="Courier-Bold",
            fontSize=16,
            leading=20,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=16,
            spaceAfter=8,
            keepWithNext=True,
        ),
        "h2": ParagraphStyle(
            "ManualH2",
            parent=base["Heading2"],
            fontName="Courier-Bold",
            fontSize=13,
            leading=16,
            textColor=colors.HexColor("#2E74B5"),
            spaceBefore=12,
            spaceAfter=6,
            keepWithNext=True,
        ),
        "h3": ParagraphStyle(
            "ManualH3",
            parent=base["Heading3"],
            fontName="Courier-Bold",
            fontSize=11,
            leading=14,
            textColor=colors.HexColor("#1F4D78"),
            spaceBefore=8,
            spaceAfter=4,
            keepWithNext=True,
        ),
        "body": ParagraphStyle(
            "ManualBody",
            parent=base["BodyText"],
            fontName="Courier",
            fontSize=9,
            leading=12,
            spaceAfter=5,
        ),
        "meta": ParagraphStyle(
            "ManualMeta",
            parent=base["BodyText"],
            fontName="Courier",
            fontSize=8,
            leading=10,
            textColor=colors.HexColor("#555555"),
            spaceAfter=3,
        ),
        "bullet": ParagraphStyle(
            "ManualBullet",
            parent=base["BodyText"],
            fontName="Courier",
            fontSize=9,
            leading=12,
            leftIndent=14,
            firstLineIndent=-8,
            spaceAfter=3,
        ),
        "code": ParagraphStyle(
            "ManualCode",
            parent=base["Code"],
            fontName="Courier",
            fontSize=7,
            leading=8.3,
            leftIndent=10,
            rightIndent=4,
            spaceAfter=2,
            textColor=colors.HexColor("#222222"),
        ),
    }


def pdf_header_footer(canvas, doc) -> None:
    canvas.saveState()
    canvas.setFont("Courier", 7)
    canvas.setFillColor(colors.HexColor("#666666"))
    canvas.drawString(inch, 0.45 * inch, "The Latticra Reference Manual")
    canvas.drawRightString(7.5 * inch, 0.45 * inch, f"Page {doc.page}")
    canvas.restoreState()


def build_pdf(docs: Sequence[SourceDoc], output_path: Path, generated_at: str) -> None:
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=letter,
        rightMargin=0.8 * inch,
        leftMargin=0.8 * inch,
        topMargin=0.8 * inch,
        bottomMargin=0.75 * inch,
        title="The Latticra Reference Manual",
        author="Latticra",
    )
    story = []
    story.append(Paragraph("The Latticra Reference Manual", styles["title"]))
    story.append(Paragraph("v0.1.0 Documentation Assembly", styles["subtitle"]))
    story.append(Paragraph(f"Generated: {generated_at}", styles["subtitle"]))

    summary_data = [
        ["Source documents", str(len(docs))],
        ["Source lines", str(sum(source.lines for source in docs))],
        ["Source words", str(sum(source.words for source in docs))],
        ["Editable source", "the-latticra-reference-manual-v0.1.0.md"],
        ["Root README link", "Top handbook links + Start Here + Main Documentation"],
    ]
    table = Table(summary_data, colWidths=[2.1 * inch, 4.3 * inch])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#E8EEF5")),
                ("TEXTCOLOR", (0, 0), (-1, -1), colors.HexColor("#111111")),
                ("FONT", (0, 0), (0, -1), "Courier-Bold", 9),
                ("FONT", (1, 0), (1, -1), "Courier", 9),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#B8C2CC")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]
        )
    )
    story.append(table)
    story.append(Spacer(1, 12))
    story.append(
        Paragraph(
            "This generated reference book preserves the current working-tree documentation corpus in a reader-oriented order. "
            "Public HTML pages and existing companion PDFs/DOCX files are indexed separately in the manifest.",
            styles["body"],
        )
    )
    story.append(PageBreak())

    docs_by_part = group_by_part(docs)
    story.append(Paragraph("Manual Part Order", styles["h1"]))
    for part in PART_ORDER:
        count = len(docs_by_part.get(part, []))
        if count:
            story.append(Paragraph(f"- {html.escape(part)}: {count} source records", styles["bullet"]))
    story.append(PageBreak())

    for part in PART_ORDER:
        part_docs = docs_by_part.get(part, [])
        if not part_docs:
            continue
        story.append(Paragraph(html.escape(part), styles["h1"]))
        for source in part_docs:
            story.append(Paragraph(f"- <font name=\"Courier\">{html.escape(source.path)}</font> - {html.escape(source.title)}", styles["bullet"]))
        story.append(PageBreak())
        for source in part_docs:
            story.append(Paragraph(html.escape(source.path), styles["h2"]))
            story.append(Paragraph(f"Source title: {html.escape(source.title)}", styles["meta"]))
            story.append(Paragraph(f"Lines: {source.lines} | Words: {source.words} | SHA-256: {source.sha256}", styles["meta"]))
            add_markdown_to_pdf_story(story, source.content, styles)
    doc.build(story, onFirstPage=pdf_header_footer, onLaterPages=pdf_header_footer)


def flush_paragraph_pdf(story: list, lines: list[str], styles: dict[str, ParagraphStyle]) -> None:
    if not lines:
        return
    text = " ".join(line.strip() for line in lines if line.strip())
    if text:
        story.append(Paragraph(html.escape(clean_inline(text)), styles["body"]))
    lines.clear()


def add_code_lines_pdf(story: list, lines: Sequence[str], styles: dict[str, ParagraphStyle]) -> None:
    if not lines:
        return
    wrapped: list[str] = []
    for line in lines:
        chunks = wrap_preserving_words(line, 96) or [""]
        wrapped.extend(chunks)
    code_text = "\n".join(html.escape(line) for line in wrapped)
    story.append(Preformatted(code_text, styles["code"], maxLineLength=96))


def add_markdown_to_pdf_story(story: list, markdown_text: str, styles: dict[str, ParagraphStyle]) -> None:
    paragraph_lines: list[str] = []
    code_lines: list[str] = []
    in_code = False
    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip("\n")
        stripped = line.strip()
        if stripped.startswith(("```", "~~~")):
            flush_paragraph_pdf(story, paragraph_lines, styles)
            if in_code:
                add_code_lines_pdf(story, code_lines, styles)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not stripped:
            flush_paragraph_pdf(story, paragraph_lines, styles)
            continue
        if is_markdown_table_line(stripped):
            flush_paragraph_pdf(story, paragraph_lines, styles)
            add_code_lines_pdf(story, [line], styles)
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            flush_paragraph_pdf(story, paragraph_lines, styles)
            level = min(3, len(heading.group(1)) + 1)
            style_name = {1: "h1", 2: "h2", 3: "h3"}[level]
            story.append(Paragraph(html.escape(clean_inline(heading.group(2))), styles[style_name]))
            continue
        bullet = re.match(r"^[-*+]\s+(.+)$", stripped)
        if bullet:
            flush_paragraph_pdf(story, paragraph_lines, styles)
            story.append(Paragraph("- " + html.escape(clean_inline(bullet.group(1))), styles["bullet"]))
            continue
        number = re.match(r"^(\d+[.)])\s+(.+)$", stripped)
        if number:
            flush_paragraph_pdf(story, paragraph_lines, styles)
            story.append(Paragraph(f"{number.group(1)} {html.escape(clean_inline(number.group(2)))}", styles["bullet"]))
            continue
        if stripped.startswith(">"):
            flush_paragraph_pdf(story, paragraph_lines, styles)
            story.append(Paragraph(html.escape(clean_inline(stripped.lstrip("> "))), styles["body"]))
            continue
        if stripped.startswith("|"):
            flush_paragraph_pdf(story, paragraph_lines, styles)
            add_code_lines_pdf(story, [line], styles)
            continue
        paragraph_lines.append(line)
    flush_paragraph_pdf(story, paragraph_lines, styles)
    if code_lines:
        add_code_lines_pdf(story, code_lines, styles)


def write_readme(
    output_dir: Path,
    docs: Sequence[SourceDoc],
    generated_at: str,
    include_docx: bool,
) -> None:
    downloads = ["- [PDF edition](the-latticra-reference-manual-v0.1.0.pdf)"]
    if include_docx:
        downloads.append("- [Editable DOCX edition](the-latticra-reference-manual-v0.1.0.docx)")
    downloads.extend(
        [
            "- [Editable Markdown source](the-latticra-reference-manual-v0.1.0.md)",
            "- [Source manifest](source-manifest.json)",
        ]
    )
    downloads_text = "\n".join(downloads)
    readme = f"""# The Latticra Reference Manual

Status: generated v0.1.0 documentation assembly
Generated: {generated_at}
Scope: canonical documentation sources in the current working tree.

## Downloads

{downloads_text}

## Corpus

- Source documents: {len(docs)}
- Source lines: {sum(doc.lines for doc in docs)}
- Source words: {sum(doc.words for doc in docs)}

## Root README Link

The root README links to this manual package from the top handbook links, Start Here, and Main Documentation.
"""
    (output_dir / "README.md").write_text(readme, encoding="utf-8")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build The Latticra Reference Manual.")
    parser.add_argument(
        "--root",
        default=Path(__file__).resolve().parents[1],
        type=Path,
        help="Repository root.",
    )
    parser.add_argument(
        "--output-dir",
        default=None,
        type=Path,
        help="Output directory. Defaults to docs/latticra-reference-manual.",
    )
    parser.add_argument(
        "--emit-docx",
        action="store_true",
        help="Also build an editable DOCX. This is opt-in because the full corpus is large.",
    )
    parser.add_argument(
        "--skip-docx",
        action="store_true",
        help=argparse.SUPPRESS,
    )
    parser.add_argument(
        "--skip-pdf",
        action="store_true",
        help="Skip the PDF build.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    root = args.root.resolve()
    output_dir = args.output_dir or (root / "docs" / "latticra-reference-manual")
    output_dir.mkdir(parents=True, exist_ok=True)

    generated_at = datetime.now().astimezone().strftime("%Y-%m-%d %H:%M:%S %Z")
    docs = discover_sources(root)
    artifacts = companion_artifacts(root)

    markdown_path = output_dir / "the-latticra-reference-manual-v0.1.0.md"
    docx_path = output_dir / "the-latticra-reference-manual-v0.1.0.docx"
    pdf_path = output_dir / "the-latticra-reference-manual-v0.1.0.pdf"
    manifest_path = output_dir / "source-manifest.json"

    build_markdown(docs, artifacts, markdown_path, generated_at)
    write_manifest(docs, artifacts, manifest_path, generated_at)
    include_docx = args.emit_docx and not args.skip_docx

    write_readme(output_dir, docs, generated_at, include_docx=include_docx)
    if include_docx:
        build_docx(docs, docx_path, generated_at)
    if not args.skip_pdf:
        build_pdf(docs, pdf_path, generated_at)

    print(f"Built {markdown_path}")
    if include_docx:
        print(f"Built {docx_path}")
    if not args.skip_pdf:
        print(f"Built {pdf_path}")
    print(f"Built {manifest_path}")


if __name__ == "__main__":
    main()
